from docx import Document
import os

def read_paragraphs():
    """ 문단 읽는 방법"""
    # print(os.getcwd())
    # doc = Document('../../../parser/data/bill/docx/2017817.docx')
    filename = '../data/docx/constitutional_law_1952'
    doc = Document(filename+'.docx')
    with open('../data/txt/constitutional_law_1952.txt', 'w') as f:

        for p in doc.paragraphs:
            print(p.text)
            f.write(p.text)
            f.write('\n')
            f.flush()



def read_table():
    """ table 읽는 방법"""
    doc = Document('data/bill/docx/2017834.docx')
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)


if __name__ == '__main__':
    read_paragraphs()
    # read_table()
