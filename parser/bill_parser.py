from docx import Document


def read_paragraphs():
    """ 문단 읽는 방법"""
    doc = Document('data/bill/2017834.docx')
    for p in doc.paragraphs:
        print(p)


def read_table():
    """ table 읽는 방법"""
    doc = Document('data/bill/2017834.docx')
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)


if __name__ == '__main__':
    read_paragraphs()
    read_table()
